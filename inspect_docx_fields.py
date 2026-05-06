#!/usr/bin/env python3
"""
Utility to inspect DOCX documents for placeholder fields.
It scans paragraphs and table cells for placeholders matching a regex pattern.

Usage:
  python inspect_docx_fields.py <docx_file> [pattern]
  pattern: a regex with one capturing group for the field name (default: {{(\w+)}})
"""
import sys
import re
try:
    from docx import Document
except ImportError:
    print("Missing dependency 'python-docx'. Install with: pip install python-docx")
    sys.exit(1)

def inspect_docx(path: str, pattern: str):
    doc = Document(path)
    regex = re.compile(pattern)
    placeholders = set()
    occurrences = []

    def scan_text(text, location):
        for match in regex.finditer(text):
            name = match.group(1)
            placeholders.add(name)
            occurrences.append({'location': location, 'text': text, 'placeholder': name})

    # Scan paragraphs
    for i, para in enumerate(doc.paragraphs, start=1):
        scan_text(para.text, f'paragraph {i}')

    # Scan tables
    for ti, table in enumerate(doc.tables, start=1):
        for ri, row in enumerate(table.rows, start=1):
            for ci, cell in enumerate(row.cells, start=1):
                for pi, para in enumerate(cell.paragraphs, start=1):
                    location = f'table {ti} cell({ri},{ci}) para {pi}'
                    scan_text(para.text, location)

    return placeholders, occurrences

def main():
    if len(sys.argv) < 2:
        print("Usage: python inspect_docx_fields.py <docx_file> [pattern]")
        sys.exit(1)
    path = sys.argv[1]
    pattern = sys.argv[2] if len(sys.argv) > 2 else r"\{\{(\w+)\}\}"  # default {{field}}
    try:
        placeholders, occurrences = inspect_docx(path, pattern)
        print("Found placeholders:", list(sorted(placeholders)))
        for occ in occurrences:
            print(f"[{occ['location']}] '{occ['placeholder']}' in text: {occ['text']}")
    except Exception as e:
        print(f"Error inspecting DOCX: {e}")
        sys.exit(1)

if __name__ == '__main__':
    main()