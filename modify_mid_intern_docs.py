#!/usr/bin/env python3
"""
Script to insert placeholder tokens into mid-internship DOCX templates.
It replaces static student and internship data fields with {{placeholders}}.
"""
import os
from docx import Document

# Mapping of template filenames to placeholder replacement rules
TEMPLATES_DIR = os.path.join('templates', 'mid-internship-docs')
FILES = [
    'obrigatorio_relatorio_de_atividades_empresa 2025-2_alimentos_EAD .docx',
    'obrigatorio_relatorio_de_supervisao_aluno 2025-2 Dra Dayana.docx'
]

def replace_paragraph(paragraph, text_map):
    """If paragraph starts with a key in text_map, replace its text."""
    text = paragraph.text.strip()
    for key, replacement in text_map.items():
        if text.startswith(key):
            paragraph.text = replacement
            return True
    return False

def process_file(path):
    print(f'Processing {path}')
    doc = Document(path)
    # Common replacements
    common_map = {
        'Estagiário:': 'Estagiário: {{nome}}    R.A.: {{ra}}',
        # Course is always 'Farmácia' for mid-internship docs
        'Curso:': 'Curso: Farmácia    Sem.: {{semestre}}',
        'Campus/Polo:': 'Campus/Polo: {{polo}}    Turno: {{turno}}',
        'E-mail:': 'E-mail: {{email}}    Telefone: ({{telefone_ddd}}) {{telefone_numero}}',
        # date line
        'Brasília/DF,': 'Brasília/DF, {{data_documento}}'
    }
    # File-specific: add Periodo da Avaliacao in supervision template
    supervision_map = {
        'Período da Avaliação': 'Período da Avaliação: {{start_date}} a {{end_date}}'
    }
    # Apply paragraph replacements (with special email handling per template)
    filename = os.path.basename(path)
    email_replaced = False
    for para in doc.paragraphs:
        text = para.text.strip()
        # Mid-internship company report: first E-mail is student, second is university
        if filename.startswith('obrigatorio_relatorio_de_atividades_empresa') and text.startswith('E-mail:'):
            if not email_replaced:
                para.text = common_map['E-mail:']
                email_replaced = True
            else:
                para.text = 'E-mail: FISCALPAULISTA@UNIP.BR        Telefone: (61) 2192-7080'
            continue
        # Mid-internship supervision report: same email logic
        if filename.startswith('obrigatorio_relatorio_de_supervisao_aluno') and text.startswith('E-mail:'):
            if not email_replaced:
                para.text = common_map['E-mail:']
                email_replaced = True
            else:
                para.text = 'E-mail: FISCALPAULISTA@UNIP.BR        Telefone: (61) 2192-7080'
            continue
        # Other placeholder replacements
        if text.startswith('Estagiário:') or text.startswith('Curso:') or \
           text.startswith('Campus/Polo:') or text.startswith('Brasília/DF,'):
            # supervision-specific period mapping
            if text.startswith('Período da Avaliação') and filename.startswith('obrigatorio_relatorio_de_supervisao_aluno'):
                replace_paragraph(para, supervision_map)
            else:
                replace_paragraph(para, common_map)
    # Fix typo 'Atividades Realizadass' in company report
    if filename.startswith('obrigatorio_relatorio_de_atividades_empresa'):
        for para in doc.paragraphs:
            if para.text.strip() == 'Atividades Realizadass':
                para.text = 'Atividades Realizadas'
    # Table replacements: first table cell(0,0)
    if doc.tables:
        tbl = doc.tables[0]
        # Row 0, cell 0: replace period placeholder
        cell = tbl.rows[0].cells[0]
        if 'PERÍODO:' in cell.text:
            cell.text = 'PERÍODO: {{start_date}} a {{end_date}}'
    # Save back
    doc.save(path)

def main():
    for filename in FILES:
        full_path = os.path.join(TEMPLATES_DIR, filename)
        if os.path.exists(full_path):
            process_file(full_path)
        else:
            print(f'File not found: {full_path}')

if __name__ == '__main__':
    main()