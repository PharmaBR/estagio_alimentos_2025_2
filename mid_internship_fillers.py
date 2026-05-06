"""
Module to fill mid-internship DOCX templates by replacing placeholders with actual data.
"""
import os
import json
import shutil
import subprocess
import fitz
import random
from docx import Document

from models import DocumentData

import os
# Paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_CONFIG_FILE = os.path.join(BASE_DIR, "internship_template.json")
TEMPLATES_DIR = os.path.join(BASE_DIR, "templates", "mid-internship-docs")
OUTPUT_DIR = os.path.join(BASE_DIR, "filled_docs")
# Signature path (absolute)
SIGNATURE_PATH = os.path.join(BASE_DIR, "templates", "assinatura_breno.png")
SIG_WIDTH_PTS = 300  # signature width in PDF points (increased for visibility)
JITTER_X = 5        # max horizontal jitter in points
JITTER_Y = 3        # max vertical jitter in points

def _format_date(iso_date: str) -> str:
    """Converts ISO date YYYY-MM-DD to DD/MM/YYYY"""
    try:
        y, m, d = iso_date.split('-')
        return f"{d}/{m}/{y}"
    except Exception:
        return iso_date

class BaseDocxFiller:
    """Base class for DOCX template fillers"""
    def __init__(self, document_data: DocumentData):
        self.data = document_data
        # Load internship template config for dates
        with open(TEMPLATE_CONFIG_FILE, 'r', encoding='utf-8') as f:
            self.config = json.load(f)
        # Label to search for signature placement in PDF
        self.signature_label = getattr(self, 'SIGNATURE_LABEL', None)

    def _get_common_mapping(self) -> dict[str, str]:
        """Returns common placeholder-to-value mapping"""
        cfg = self.config
        user = self.data.user
        return {
            'nome': user.nome,
            'ra': user.ra,
            'disciplina_estagio': self.data.internship.disciplina_estagio,
            'semestre': user.semestre,
            'polo': user.polo,
            'turno': f"{cfg.get('start_time')} às {cfg.get('end_time')}",
            'email': user.email,
            'telefone_ddd': user.telefone_ddd,
            'telefone_numero': user.telefone_numero,
            'data_documento': cfg.get('document_date', ''),
            'start_date': _format_date(cfg.get('start_date', '')),
            'end_date': _format_date(cfg.get('end_date', '')),
        }

    def _replace_placeholders(self, doc: Document, mapping: dict[str, str]):
        """Replaces all placeholders in paragraphs and table cells"""
        for para in doc.paragraphs:
            for key, val in mapping.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, val)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, val in mapping.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, val)

    def _ensure_output_dir(self):
        os.makedirs(OUTPUT_DIR, exist_ok=True)

    def save_doc(self, doc: Document, output_name: str) -> str:
        """Saves the filled Document to OUTPUT_DIR and returns its path"""
        self._ensure_output_dir()
        output_path = os.path.join(OUTPUT_DIR, output_name)
        doc.save(output_path)
        return output_path

    def _convert_to_pdf(self, docx_path: str) -> str:
        """Attempts to convert a DOCX file to PDF using LibreOffice."""
        base = os.path.splitext(os.path.basename(docx_path))[0]
        pdf_path = os.path.join(OUTPUT_DIR, f"{base}.pdf")
        # Use LibreOffice if available
        if shutil.which('soffice'):
            cmd = [
                'soffice', '--headless',
                '--convert-to', 'pdf',
                '--outdir', OUTPUT_DIR,
                docx_path
            ]
            subprocess.run(cmd, check=True)
            # After converting DOCX to PDF, overlay signature (and name/CRF)
            self._add_signature_to_pdf(pdf_path)
            return pdf_path
        # No converter available, return DOCX path
        return docx_path
    
    def _add_signature_to_pdf(self, pdf_path: str):
        """Overlays the signature image over the 'INSTITUIÇÃO DE ENSINO' text in the PDF."""
        # Only overlay signature for company docs
        label = getattr(self, 'signature_label', None)
        if not label or not os.path.exists(SIGNATURE_PATH):
            return
        try:
            doc = fitz.open(pdf_path)
            # Find page containing label
            target_page = None
            for p in doc:
                if p.search_for(label):
                    target_page = p
                    break
            if target_page is None:
                target_page = doc[-1]
            # Load image and compute size
            pix = fitz.Pixmap(SIGNATURE_PATH)
            scale = SIG_WIDTH_PTS / pix.width
            sig_h = pix.height * scale
            # Determine placement
            instances = target_page.search_for(label)
            if instances:
                rect = instances[0]
                # Horizontally center signature on label
                x = rect.x0 + (rect.x1 - rect.x0) / 2 - SIG_WIDTH_PTS / 2
                # Vertically center over the label text area to overlap it
                y = rect.y0 + ((rect.y1 - rect.y0) - sig_h) / 2
            else:
                page_rect = target_page.rect
                x = (page_rect.width - SIG_WIDTH_PTS) / 2
                y = page_rect.height - sig_h - 80
            # Jitter
            x += random.uniform(-JITTER_X, JITTER_X)
            y += random.uniform(-JITTER_Y, JITTER_Y)
            # Insert signature image
            sig_rect = fitz.Rect(x, y, x + SIG_WIDTH_PTS, y + sig_h)
            target_page.insert_image(sig_rect, filename=SIGNATURE_PATH, overlay=True)
            # Add signer name and CRF under signature
            text_y = y + sig_h + 4
            target_page.insert_text((x, text_y), "Prof. Breno Silva de Abreu", fontsize=12)
            target_page.insert_text((x, text_y + 14), "CRF-DF 2173", fontsize=12)
            doc.save(pdf_path)
            doc.close()
        except Exception:
            pass


class CompanyActivitiesDocxFiller(BaseDocxFiller):
    """Fills the company activities mid-internship report"""
    TEMPLATE_NAME = 'obrigatorio_relatorio_de_atividades_empresa 2025-2_alimentos_EAD .docx'
    # PDF signature should overlay next to this label
    SIGNATURE_LABEL = 'INSTITUIÇÃO DE ENSINO'

    def fill(self) -> str:
        template_path = os.path.join(TEMPLATES_DIR, self.TEMPLATE_NAME)
        doc = Document(template_path)
        mapping = self._get_common_mapping()
        # Override document date for mid-internship
        mapping['data_documento'] = '03 de Abril de 2026'
        self._replace_placeholders(doc, mapping)
        # Save DOCX
        docx_path = self.save_doc(doc, f"relatorio_atividades_empresa_{self.data.user.ra}.docx")
        return self._convert_to_pdf(docx_path)


class SupervisionReportDocxFiller(BaseDocxFiller):
    """Fills the student supervision mid-internship report"""
    TEMPLATE_NAME = 'obrigatorio_relatorio_de_supervisao_aluno 2025-2 Dra Dayana.docx'
    # No signature for supervision report
    SIGNATURE_LABEL = None

    def fill(self) -> str:
        template_path = os.path.join(TEMPLATES_DIR, self.TEMPLATE_NAME)
        doc = Document(template_path)
        mapping = self._get_common_mapping()
        # Override document date for mid-internship
        mapping['data_documento'] = '03 de Abril de 2026'
        self._replace_placeholders(doc, mapping)
        # Save DOCX
        docx_path = self.save_doc(doc, f"relatorio_supervisao_aluno_{self.data.user.ra}.docx")
        # Tentar converter para PDF, senão retornar DOCX
        try:
            return self._convert_to_pdf(docx_path)
        except RuntimeError:
            return docx_path